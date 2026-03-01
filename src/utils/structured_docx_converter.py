"""
DOCX Converter - Converts structured TailoredResume (Pydantic) to a .docx file
Equivalent to StructuredLaTeXConverter but outputs Word format for better ATS parsing.

Requires: pip install python-docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.config.constants import ValidLanguages
from src.config.schemas import PersonalInfo


def _add_horizontal_rule(paragraph):
    """Add a bottom border to a paragraph to simulate a section divider."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_border(cell, **kwargs):
    """Utility to remove table cell borders — used for layout tables."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


class StructuredDocxConverter:
    """
    Converts a TailoredResume Pydantic object into a .docx Word document.
    Mirrors the structure of StructuredLaTeXConverter.

    Key design decisions for ATS compatibility:
    - No tables for layout (ATS parsers often skip table content)
    - No headers/footers (ATS parsers often skip them)
    - Clean semantic structure: name → headline → summary → experience → skills
    - Standard fonts only (Calibri)
    - All contact info in the body, not in a header
    """

    # ── Styling constants ─────────────────────────────────────────

    FONT_NAME        = "Calibri"
    COLOR_BLACK      = RGBColor(0, 0, 0)
    COLOR_DARK_GRAY  = RGBColor(64, 64, 64)

    SIZE_NAME        = Pt(20)
    SIZE_HEADLINE    = Pt(11)
    SIZE_CONTACT     = Pt(10)
    SIZE_SECTION     = Pt(12)
    SIZE_BODY        = Pt(10)
    SIZE_JOB_TITLE   = Pt(10)

    def __init__(self, language: ValidLanguages = ValidLanguages.EN):
        self.english = (language == ValidLanguages.EN)

    # ── Helpers ───────────────────────────────────────────────────

    def _set_font(self, run, size=None, bold=False, italic=False, color=None):
        run.font.name = self.FONT_NAME
        if size:
            run.font.size = size
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def _add_paragraph(self, doc, text="", alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=0, space_after=4) -> object:
        p = doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            self._set_font(run, size=self.SIZE_BODY)
        return p

    def _add_section_heading(self, doc, title: str):
        """Bold heading with a bottom border rule — mirrors LaTeX \section*{}[\titlerule]."""
        p = self._add_paragraph(doc, space_before=8, space_after=2)
        run = p.add_run(title.upper())
        self._set_font(run, size=self.SIZE_SECTION, bold=True)
        _add_horizontal_rule(p)

    # ── Section builders ──────────────────────────────────────────

    def _build_header(self, doc, name: str, headline: str, personal_info: PersonalInfo):
        """Name, headline, contact info — all in body (never in Word header/footer)."""

        # Name
        p = self._add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        run = p.add_run(name)
        self._set_font(run, size=self.SIZE_NAME, bold=True)

        # Headline
        p = self._add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        run = p.add_run(headline)
        self._set_font(run, size=self.SIZE_HEADLINE, italic=True)

        # Contact line
        info = personal_info or PersonalInfo()
        parts = []
        if info.email:    parts.append(info.email)
        if info.location: parts.append(info.location)
        if info.linkedin: parts.append(f"linkedin.com/in/{info.linkedin}")
        if info.github:   parts.append(f"github.com/{info.github}")

        if parts:
            p = self._add_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
            run = p.add_run(" | ".join(parts))
            self._set_font(run, size=self.SIZE_CONTACT, color=self.COLOR_DARK_GRAY)

    def _build_summary(self, doc, summary: str):
        self._add_section_heading(doc, "Profile" if self.english else "Perfil")
        p = self._add_paragraph(doc, space_after=6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(summary)
        self._set_font(run, size=self.SIZE_BODY)

    def _build_experience(self, doc, experience: list):
        self._add_section_heading(doc, "Professional Experience" if self.english else "Experiencia")

        for entry in experience:
            p = self._add_paragraph(doc, space_before=4, space_after=0)
            run = p.add_run(entry.job_title)
            self._set_font(run, size=self.SIZE_BODY, bold=True)

            p = self._add_paragraph(doc, space_before=0, space_after=0)
            run = p.add_run(f"{entry.company}")
            self._set_font(run, size=self.SIZE_JOB_TITLE, bold=True, italic=True)
            run = p.add_run(f"  {suffix}" if suffix else None)
            self._set_font(run, size=self.SIZE_BODY, italic=True)

            p = self._add_paragraph(doc, space_before=0, space_after=0)
            run = p.add_run(entry.location)
            self._set_font(run, size=self.SIZE_BODY, italic=True)

            p = self._add_paragraph(doc, space_before=0, space_after=2)
            run = p.add_run(f"{entry.start_date} – {entry.end_date}")
            self._set_font(run, size=self.SIZE_BODY)
            #p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Bullet tasks
            for task in entry.tasks:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(task)
                self._set_font(run, size=self.SIZE_BODY)

            # Spacing between entries
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _build_skills(self, doc, skills: list):
        self._add_section_heading(doc, "Skills")

        for entry in skills:
            p = self._add_paragraph(doc, space_before=2, space_after=2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Category bold, skills normal
            run = p.add_run(f"{entry.category_name}: ")
            self._set_font(run, size=self.SIZE_BODY, bold=True)
            run = p.add_run(", ".join(entry.skills))
            self._set_font(run, size=self.SIZE_BODY)

    def _build_education(self, doc):
        self._add_section_heading(doc, "Education")

        entries = [
            {
                "degree": "Master in Smart Systems",
                "gpa": "GPA: 9.0/10",
                "institution": "University of Salamanca",
                "location": "Spain",
                "dates": "11/2023 – 10/2024",
                "bullets": [
                    "Specialization: Machine Learning, Deep Learning, Statistics, Data Mining, Data Science, Data Visualization, Econometrics",
                    'Thesis: "Application of Convolutional Neural Networks in bioacoustics analysis" - Applied advanced analytical modeling techniques to solve complex pattern recognition problem',
                ],
            },
            {
                "degree": "Bachelor of Science in Electronic Engineering",
                "gpa": "GPA: 8.9/10",
                "institution": "Catholic University of Argentina (UCA)",
                "location": "Argentina",
                "dates": "03/2016 – 07/2022",
                "bullets": [
                    "Quantitative curriculum including Applied Mathematics, Operations Research, Statistics, Systems Engineering",
                    'Thesis: "Image stabilizer system - application of the Kalman Filter using dual-core DSP"',
                ],
            },
        ]

        for edu in entries:
            p = self._add_paragraph(doc, space_before=4, space_after=0)
            run = p.add_run(edu["degree"])
            self._set_font(run, size=self.SIZE_BODY, bold=True)
            run = p.add_run(f"  | {edu['gpa']}")
            self._set_font(run, size=self.SIZE_BODY,)

            p = self._add_paragraph(doc, space_before=0, space_after=0)
            run = p.add_run(edu['institution'])
            self._set_font(run, size=self.SIZE_BODY, bold=True)

            p = self._add_paragraph(doc, space_before=0, space_after=0)
            run = p.add_run(f"{edu['location']}")
            self._set_font(run, size=self.SIZE_BODY, italic=True)

            p = self._add_paragraph(doc, space_before=0, space_after=2)
            run = p.add_run(edu["dates"])
            self._set_font(run, size=self.SIZE_BODY)

            for bullet in edu["bullets"]:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(bullet)
                self._set_font(run, size=self.SIZE_BODY)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _build_certifications(self, doc):
        self._add_section_heading(doc, "Certifications & Professional Development")

        certs = [
            "AWS Cloud Practitioner (2025)",
            "Generative AI with Large Language Models Specialization - DeepLearning.AI & AWS (2025)",
            "Neural Networks Certification (2023)",
            "Machine Learning Certification (2022)",
            "MIT Leading Digital Transformation - Santander Scholarship (2020)",
            "PMP Candidate - Project Management Professional (in progress)",
        ]

        for cert in certs:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(cert)
            self._set_font(run, size=self.SIZE_BODY)

    def _build_awards(self, doc):
        self._add_section_heading(doc, "Awards & Recognition")

        awards = [
            "Winner - Argentine Association of Control's National Thesis Contest (2023)",
            "Recognition from Microchip Technology and MC Electronics for thesis excellence",
            "PRIUNES Scholarship - Catholic University of Argentina (2016)",
        ]

        for award in awards:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(award)
            self._set_font(run, size=self.SIZE_BODY)

    # ── Main entry point ──────────────────────────────────────────

    def convert(self, resume, personal_info: PersonalInfo = None) -> Document:
        """
        Convert a TailoredResume Pydantic object to a python-docx Document.

        Returns the Document object — call .save(path) to write to disk.
        Unlike StructuredLaTeXConverter which returns a string, we return
        the Document directly because .docx is binary and can't be passed
        around as a string.
        """
        doc = Document()

        # ── Page margins (0.75in all sides, matching LaTeX converter) ──
        for section in doc.sections:
            section.top_margin    = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin   = Inches(0.75)
            section.right_margin  = Inches(0.75)

        # ── Remove default paragraph spacing ──
        style = doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = self.SIZE_BODY
        style.paragraph_format.space_after = Pt(0)

        info = personal_info or PersonalInfo()

        self._build_header(doc,
                           name=info.name or "Your Name",
                           headline=resume.headline,
                           personal_info=info)
        self._build_summary(doc, resume.professional_summary)
        self._build_experience(doc, resume.professional_experience)
        self._build_education(doc)
        self._build_certifications(doc)
        self._build_awards(doc)
        self._build_skills(doc, resume.skills)

        return doc